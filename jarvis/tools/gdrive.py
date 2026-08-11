"""Google Drive tools: search, read, and save documents."""

from __future__ import annotations

import io
from pathlib import Path

from anthropic import beta_tool

from . import ToolContext, as_document, cancelled_by_user
from .local_files import PathNotAllowed, resolve_safe

_MAX_EXPORT_BYTES = 200_000

# Google-native docs are exported to plain text; regular text files download as-is.
_EXPORTABLE = {
    "application/vnd.google-apps.document": "text/plain",
    # text/csv returns only the FIRST worksheet, with no hint that the
    # others exist. Exporting as xlsx keeps every sheet; it is parsed
    # below. Needs no extra scope.
    "application/vnd.google-apps.spreadsheet":
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.google-apps.presentation": "text/plain",
}


def _drive(ctx: ToolContext):
    return ctx.google_service("drive", "v3")


# Binary formats Drive will not convert for us, read locally instead. These
# are the bulk of a real company Drive, and drive_read used to refuse them.
_BINARY_READABLE = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

_XLSX = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _extract_text(raw: bytes, mime: str, name: str) -> str:
    """Turn downloaded bytes into readable text.

    Import failures degrade to an explanation rather than an exception: the
    document extras are optional, and a missing one must not look like a
    broken file.
    """
    import io as _io

    try:
        if mime == "application/pdf":
            from pypdf import PdfReader

            reader = PdfReader(_io.BytesIO(raw))
            pages = [(pg.extract_text() or "").strip() for pg in reader.pages]
            body = "\n\n".join(
                f"[page {i + 1}]\n{t}" for i, t in enumerate(pages) if t
            )
            return body or ("(no extractable text — this PDF is probably a scan, "
                            "so it would need OCR)")
        if mime == _DOCX:
            import docx

            d = docx.Document(_io.BytesIO(raw))
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for table in d.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        if mime == _XLSX:
            import openpyxl

            wb = openpyxl.load_workbook(_io.BytesIO(raw), read_only=True, data_only=True)
            out = []
            for ws in wb.worksheets:
                out.append(f"[sheet: {ws.title}]")
                for row in ws.iter_rows(values_only=True):
                    if any(v is not None and str(v).strip() for v in row):
                        out.append(" | ".join("" if v is None else str(v) for v in row))
            return "\n".join(out)
    except ImportError:
        return (f"'{name}' needs the document extras to read. Install them with "
                "pip install .[documents]")
    except Exception as e:
        return f"'{name}' could not be parsed ({type(e).__name__}: {e})."
    return raw.decode("utf-8", errors="replace")


def build_tools(ctx: ToolContext) -> list:
    @beta_tool
    def drive_search(query: str, max_results: int = 10) -> str:
        """Search the user's Google Drive for files by name or content keywords.

        Args:
            query: What to look for, e.g. a document title or topic keywords.
            max_results: Maximum number of results to return (1-25).
        """
        if not ctx.permissions.require("drive_read", "search and read your Google Drive"):
            return "The user declined Google Drive access."
        try:
            safe_q = query.replace("\\", "\\\\").replace("'", "\\'")
            resp = _drive(ctx).files().list(
                q=f"(name contains '{safe_q}' or fullText contains '{safe_q}') and trashed = false",
                pageSize=max(1, min(int(max_results), 25)),
                fields="files(id, name, mimeType, modifiedTime, owners(displayName))",
                # Team documents live on Shared Drives, and files().list()
                # excludes those by default — so the files people most often
                # ask for were reported as not existing.
                includeItemsFromAllDrives=True,
                supportsAllDrives=True,
                corpora="allDrives",
            ).execute()
            files = resp.get("files", [])
            ctx.audit.record("tool_call", tool="drive_search", detail=query)
            if not files:
                return f"No Drive files matched '{query}'."
            lines = [
                f"- {f['name']}  (id: {f['id']}, type: {f['mimeType'].split('.')[-1]}, "
                f"modified: {f.get('modifiedTime', '?')[:10]})"
                for f in files
            ]
            # Drive titles are attacker-influenceable content — wrap as data.
            return as_document(f"drive-search:{query}", "\n".join(lines))
        except Exception as e:  # googleapiclient raises many transport error types
            ctx.audit.record("tool_call", tool="drive_search", detail=query, ok=False)
            return f"Drive search failed: {e}"

    @beta_tool
    def drive_read(file_id: str) -> str:
        """Read the text content of a Google Drive file (Docs, Sheets, Slides,
        or plain-text files). Use drive_search first to find the file id.

        Args:
            file_id: The Drive file id to read.
        """
        if not ctx.permissions.require("drive_read", "search and read your Google Drive"):
            return "The user declined Google Drive access."
        try:
            from googleapiclient.http import MediaIoBaseDownload

            service = _drive(ctx)
            meta = service.files().get(fileId=file_id, fields="id, name, mimeType, size").execute()
            mime = meta["mimeType"]
            name = meta["name"]

            if mime in _EXPORTABLE:
                request = service.files().export_media(fileId=file_id, mimeType=_EXPORTABLE[mime])
            elif mime.startswith("text/") or mime in ("application/json",):
                request = service.files().get_media(fileId=file_id)
            elif mime in _BINARY_READABLE:
                request = service.files().get_media(fileId=file_id)
            else:
                return (
                    f"'{name}' is {mime}, which I can't read as text. I can read Google "
                    "Docs, Sheets and Slides, PDF, Word and Excel files, and plain text."
                )

            buf = io.BytesIO()
            # Small chunks so the byte cap actually bounds the transfer
            # (default chunk size is 100 MB, which would defeat the loop guard).
            downloader = MediaIoBaseDownload(buf, request, chunksize=64 * 1024)
            done = False
            while not done and buf.tell() <= _MAX_EXPORT_BYTES:
                _, done = downloader.next_chunk()
            truncated = (not done) or buf.tell() > _MAX_EXPORT_BYTES
            raw = buf.getvalue()[:_MAX_EXPORT_BYTES]
            effective = _EXPORTABLE.get(mime, mime)
            text = _extract_text(raw, effective, name)
            ctx.audit.record("tool_call", tool="drive_read", detail=f"{name} ({file_id})")
            doc = as_document(f"gdrive:{name}", text)
            return doc + ("\n[Truncated — document is larger than the read limit.]" if truncated else "")
        except Exception as e:
            ctx.audit.record("tool_call", tool="drive_read", detail=file_id, ok=False)
            return f"Drive read failed: {e}"

    @beta_tool
    def drive_save_text(name: str, content: str, folder_id: str = "") -> str:
        """Create a new text file in the user's Google Drive. Asks the user to
        confirm before creating.

        Args:
            name: File name to create, e.g. "meeting-notes.txt".
            content: Full text content of the file.
            folder_id: Optional Drive folder id to place the file in.
        """
        if not ctx.permissions.require("drive_write", "create files in your Google Drive"):
            return "The user declined Google Drive write access."
        summary = f"I will create '{name}' in your Google Drive ({len(content)} characters)."
        result = ctx.confirmer.confirm("drive_save_text", summary)
        if not result:
            return cancelled_by_user(result, "the Drive upload")
        try:
            from googleapiclient.http import MediaIoBaseUpload

            body: dict = {"name": name}
            if folder_id:
                body["parents"] = [folder_id]
            media = MediaIoBaseUpload(
                io.BytesIO(content.encode("utf-8")), mimetype="text/plain"
            )
            created = _drive(ctx).files().create(
                body=body, media_body=media, fields="id, webViewLink"
            ).execute()
            ctx.audit.record("tool_call", tool="drive_save_text", detail=name,
                             decision="confirmed")
            return f"Created '{name}' in Drive: {created.get('webViewLink', created['id'])}"
        except Exception as e:
            ctx.audit.record("tool_call", tool="drive_save_text", detail=name, ok=False)
            return f"Drive upload failed: {e}"

    @beta_tool
    def drive_upload(local_path: str, folder_id: str = "") -> str:
        """Upload a local file (from the allowed directories) to Google Drive.
        Asks the user to confirm before uploading.

        Args:
            local_path: Path of the local file to upload.
            folder_id: Optional Drive folder id to place the file in.
        """
        if not ctx.permissions.require("drive_write", "upload files to your Google Drive"):
            return "The user declined Google Drive write access."
        try:
            source = resolve_safe(ctx, local_path)
        except PathNotAllowed as e:
            ctx.audit.record("tool_call", tool="drive_upload", detail=local_path,
                             decision="blocked_path", ok=False)
            return f"Blocked: {e}"
        if not source.is_file():
            return f"'{source}' does not exist or is not a file."

        size_kb = source.stat().st_size // 1024
        summary = f"I will upload {source.name} ({size_kb} KB) from {source.parent} to your Google Drive."
        result = ctx.confirmer.confirm("drive_upload", summary)
        if not result:
            return cancelled_by_user(result, "the upload")
        try:
            from googleapiclient.http import MediaFileUpload

            body: dict = {"name": source.name}
            if folder_id:
                body["parents"] = [folder_id]
            media = MediaFileUpload(str(source), resumable=True)
            created = _drive(ctx).files().create(
                body=body, media_body=media, fields="id, webViewLink"
            ).execute()
            ctx.audit.record("tool_call", tool="drive_upload", detail=str(source),
                             decision="confirmed")
            return f"Uploaded '{source.name}': {created.get('webViewLink', created['id'])}"
        except Exception as e:
            ctx.audit.record("tool_call", tool="drive_upload", detail=str(source), ok=False)
            return f"Drive upload failed: {e}"

    return [drive_search, drive_read, drive_save_text, drive_upload]
